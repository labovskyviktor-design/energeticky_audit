import os
from docx import Document
from docx.shared import Pt
import re

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    
    for line in lines:
        line = line.strip('\n')
        
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue

        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), 2)
            
        # Handle Horizontal Rules
        elif line.strip() == '---':
            doc.add_page_break() # or just a line, but page break is cleaner for appendix
            
        # Handle Bullet Points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
            
        # Handle Table placeholders (very basic)
        elif '|' in line and '-' not in line and len(line.split('|')) > 2:
            # For now, just add as normal text or bold if it's a header
            # Proper table conversion is complex, so let's just use normal formatting
            p = doc.add_paragraph()
            run = p.add_run(line.replace('|', '  '))
            run.font.size = Pt(10)
            
        # Normal paragraphs
        elif line.strip():
            # Clean up bold/italic markers for docx
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
            # Remove link syntax [text](url) -> text (url)
            clean_line = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', clean_line)
            # Remove emoji if possible, or leave them (Word handles common ones)
            
            doc.add_paragraph(clean_line)
            
        else:
            # Empty line
            pass

    doc.save(docx_path)
    print(f"File saved to {docx_path}")

if __name__ == "__main__":
    md_file = r'c:\Users\42191\.gemini\antigravity\scratch\energy-audit\README.md'
    docx_file = r'c:\Users\42191\.gemini\antigravity\scratch\energy-audit\README_Priloha.docx'
    markdown_to_docx(md_file, docx_file)
